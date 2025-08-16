from flask import Flask, request, jsonify, send_file  
import base64
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
import re

app = Flask(__name__)

def set_rtl_paragraph(paragraph):
    """Setting right-to-left directionality for a paragraph - simple approach"""
    try:
        # Setting alignment to right
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Adding XML settings directly
        pPr = paragraph._element.get_or_add_pPr()

        # bidi - bi-directional text
        if pPr.find(qn('w:bidi')) is None:
            bidi = OxmlElement('w:bidi')
            bidi.set(qn('w:val'), '1')
            pPr.append(bidi)

        # Setting RTL direction for all runs
        for run in paragraph.runs:
            rPr = run._element.get_or_add_rPr()
            if rPr.find(qn('w:rtl')) is None:
                rtl = OxmlElement('w:rtl')
                rtl.set(qn('w:val'), '1')
                rPr.append(rtl)
                
    except Exception as e:
        print(f"RTL error: {e}")

@app.route('/', methods=['GET'])
def home():
    return "📋 Resume Agent API is up and running!", 200

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({"status": "ok"}), 200

@app.route('/extract_resume', methods=['POST'])
def extract_resume():
    data = request.json
    if not data or 'resume' not in data:
        return jsonify({"error": "Missing resume field"}), 400
    resume_b64 = data['resume']
    mime_type = data.get('resumeMimeType', '')
    try:
        file_bytes = base64.b64decode(resume_b64)
    except Exception as e:
        return jsonify({"error": f"Invalid base64: {str(e)}"}), 400
    try:
        if mime_type == "application/pdf":
            reader = PdfReader(BytesIO(file_bytes))
            text = "\n".join([page.extract_text() or "" for page in reader.pages])
        elif mime_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ]:
            doc = Document(BytesIO(file_bytes))
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            return jsonify({"error": f"Unsupported MIME type: {mime_type}"}), 415
    except Exception as e:
        return jsonify({"error": f"Error reading file: {str(e)}"}), 500
    return jsonify({
        "resumeText": text,
        "length": len(text),
        "jobLink": data.get("jobLink"),
        "companyLink": data.get("companyLink"),
        "linkedinProfile": data.get("linkedinProfile"),
        "resumeFileName": data.get("resumeFileName")
    })

@app.route('/generate_docx', methods=['POST'])
def generate_docx():
    data = request.json
    text = data.get("text", "").strip()
    
    if not text:
        return jsonify({"error": "Missing 'text' field"}), 400
    
    try:
        doc = Document()

        # Main Title
        title = doc.add_heading('📋 Interview Preparation - Gemini AI', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0, 51, 102)

        # Separator Line
        separator = doc.add_paragraph('═' * 60)
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Cleaning text
        cleaned_text = (text.replace('***', '')
                           .replace('**', '')
                           .replace('###', '')
                           .replace('##', '')
                           .replace('#', '')
                           .replace('```', '')
                           .replace('__', '')
                           .replace('*', '')
                           .replace('---', '')
                           .replace('–', '-')
                           .replace('\r', '')
                           .replace('\u200E', ''))

        # Split into paragraphs and simple, clear formatting
        sections = cleaned_text.split('\n\n')
        for section in sections:
            if section.strip():
                lines = section.strip().split('\n')
                for line in lines:
                    line = line.strip()
                    if line:
                        # Removing numbers from the beginning of the line
                        clean_line = re.sub(r'^\d+\.\s*', '', line)

                        # Improved heading detection
                        is_heading = False
                        heading_keywords = ['ניתוח', 'התאמה', 'החברה', 'שאלות', 'משפטי מפתח', 'Akamai', 'קורות החיים']

                        # Checking if the line is a heading (short and contains a keyword)
                        if (len(clean_line.split()) <= 8 and
                            any(keyword in clean_line for keyword in heading_keywords)):
                            is_heading = True

                        # Additional check for headings starting with certain phrases
                        if any(clean_line.startswith(prefix) for prefix in ['ניתוח', 'התאמה', 'שאלות', 'משפטי']):
                            is_heading = True
                        
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run(clean_line)

                        # Heading formatting
                        if is_heading:
                            run.font.bold = True
                            run.font.size = Pt(15)  # Larger for headings
                            run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
                        else:
                            run.font.bold = False
                            run.font.size = Pt(11)  # Regular text

                        # Uniform font with support for Hebrew and English
                        run.font.name = 'Calibri'

                        # Aligning text to the right
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                        # Activating RTL only after the paragraph is ready
                        set_rtl_paragraph(paragraph)

        # Creating a temporary file
        temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp.name)
        
        return send_file(temp.name,
                         as_attachment=True,
                         download_name='Interview_Preparation.docx',
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                         
    except Exception as e:
        return jsonify({"error": f"Failed to generate DOCX: {str(e)}"}), 500

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 10000))
    app.run(host='0.0.0.0', port=port)